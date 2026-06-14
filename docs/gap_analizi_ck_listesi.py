"""
Declaro — Eksik ÇK / YAML / Pipeline Gap Analizi Word belgesi oluşturur.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Yardımcı fonksiyonlar ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get(edge, "CCCCCC"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def header_row(table, cols: list[str], bg="1F3864", fg="FFFFFF"):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = text
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def data_row(table, row_idx: int, values: list[str], bg=None):
    row = table.rows[row_idx]
    default_bg = "EBF0F7" if row_idx % 2 == 0 else "FFFFFF"
    for i, text in enumerate(values):
        cell = row.cells[i]
        cell.text = text
        set_cell_bg(cell, bg or default_bg)
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.font.size = Pt(9)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text: str, level: int = 1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_info_box(doc, text: str, bg="FFF3CD", border="FFC107"):
    """Uyarı / bilgi kutusu paragrafı."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    return p


# ─── Belge oluştur ────────────────────────────────────────────────────────────

doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Varsayılan stil
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ─── Başlık ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DECLARO")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor.from_string("1F3864")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Eksik Çalışma Kağıdı / YAML / Pipeline Gap Analizi")
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor.from_string("2E75B6")

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(f"Hazırlayan: Declaro Teknik Ekibi  |  Tarih: {datetime.date.today().strftime('%d.%m.%Y')}")
run3.font.size = Pt(9)
run3.font.color.rgb = RGBColor.from_string("666666")

doc.add_paragraph()

# ─── Bölüm 1: Mevcut Durum ────────────────────────────────────────────────────
add_heading(doc, "1. Mevcut Durum Özeti", 1)

doc.add_paragraph(
    "Aşağıdaki tablo, Declaro platformunda ticari kar → mali kar hesaplama pipeline'ının "
    "hangi bölümlerinin tamamlandığını, hangilerinin eksik olduğunu özetlemektedir."
)

durum_table = doc.add_table(rows=9, cols=4)
durum_table.style = "Table Grid"
durum_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Sütun genişlikleri
for i, w in enumerate([Cm(6), Cm(3.5), Cm(3.5), Cm(3.5)]):
    for row in durum_table.rows:
        row.cells[i].width = w

header_row(durum_table, ["Bölüm / Alan", "Excel ÇK", "YAML Kalem", "Durum"])

durum_data = [
    ("Zarar Olsa Dahi (Kod 297–387)", "37 ÇK", "48 YAML", "✓ Tamamlandı"),
    ("Kazancın Bulunması Halinde (Kod 401–483)", "47 ÇK", "45 YAML", "✓ Tamamlandı"),
    ("KVK Md. 32/A — YTB İndirimli KV", "1 ÇK (docs/)", "0 YAML", "✗ EKSİK"),
    ("KVK Md. 32/7 — KOBİ İndirimi", "—", "0 YAML", "✗ EKSİK"),
    ("KVK Md. 32/8 — İhracat İndirimi", "—", "0 YAML", "✗ EKSİK"),
    ("KVK Md. 9 — Geçmiş Yıl Zararı Mahsubu", "—", "0 YAML", "✗ EKSİK"),
    ("Mahsup Edilecek Vergiler (GV, Stopaj, Yurt Dışı)", "—", "0 YAML", "✗ EKSİK"),
    ("YİAKV / KVK 32/C Özet Sayfası", "1 ÇK (docs/)", "0 YAML", "⚠ Kısmi"),
]

for i, row_data in enumerate(durum_data, start=1):
    bg = None
    if "EKSİK" in row_data[3]:
        bg = "FFE6E6"
    elif "Kısmi" in row_data[3]:
        bg = "FFF3CD"
    elif "Tamamlandı" in row_data[3]:
        bg = "E6F4EA"
    data_row(durum_table, i, list(row_data), bg=bg)

doc.add_paragraph()

# ─── Bölüm 2: GRUP A ──────────────────────────────────────────────────────────
add_heading(doc, "2. Grup A — Kritik Pipeline Boşlukları", 1)
doc.add_paragraph(
    "Bu gruptaki kalemler, pipeline hesaplama adımlarında şu an 'v1'de atlanıyor / uygulanmıyor' "
    "notu ile geçilen aşamalara karşılık gelmektedir (Adım 5, 8 ve 12). "
    "En acil önceliktir; doğru vergi hesabı için zorunludur."
)

grp_a_table = doc.add_table(rows=9, cols=6)
grp_a_table.style = "Table Grid"
grp_a_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate([Cm(1.5), Cm(4.5), Cm(3), Cm(2), Cm(2), Cm(3)]):
    for row in grp_a_table.rows:
        row.cells[i].width = w

header_row(grp_a_table, ["Kod", "ÇK Adı", "Mevzuat", "Pipeline Adımı", "Beyanname Satırı", "Notlar"])

grp_a_data = [
    ("ÇK-A1", "Geçmiş Yıl Zararları Mahsubu", "KVK Md. 9", "Adım 5", "580", "Son 5 yıl; her yıl ayrı satır. En sık kullanılan kalem."),
    ("ÇK-A2", "İndirimli KV — Tek YTB Senaryosu", "KVK Md. 32/A", "Adım 8", "710 / 720", "Matrah bölüştürme; indirimli oran; kalan matrah hesabı."),
    ("ÇK-A3", "İndirimli KV — Çok YTB Senaryosu", "KVK Md. 32/A", "Adım 8", "710 / 720", "docs/KVK_32A_Cok_YTB_Calisma_Kagidi_1.xlsx MEVCUT; sadece YAML/pipeline gerekiyor."),
    ("ÇK-A4", "KOBİ İndirimli KV (%5 puan)", "KVK Md. 32/7", "Adım 8", "730", "Sanayi siciline kayıtlı üretim işletmeleri. 2022+ yürürlükte."),
    ("ÇK-A5", "İhracat İndirimli KV (%5 puan)", "KVK Md. 32/8", "Adım 8", "740", "İhracat kazancı / toplam kazanç oranı ile matrah bölüştürme."),
    ("ÇK-A6", "Geçici Vergi Mahsubu", "KVK Md. 44", "Adım 12", "640", "4 dönem geçici vergi toplamı."),
    ("ÇK-A7", "Yurt Dışında Ödenen Vergi Mahsubu", "KVK Md. 33", "Adım 12", "650", "Ülke bazlı; ÇVÖA kapsamı kontrolü."),
    ("ÇK-A8", "Tevkifat (Stopaj) Mahsubu", "KVK Md. 34", "Adım 12", "660", "Kira, menkul kıymet stopajları."),
]

for i, row_data in enumerate(grp_a_data, start=1):
    data_row(grp_a_table, i, list(row_data))

doc.add_paragraph()

# ─── Bölüm 3: GRUP B ──────────────────────────────────────────────────────────
add_heading(doc, "3. Grup B — Kazancın Bulunması Halinde Serisi Eksikleri", 1)
doc.add_paragraph(
    "Kazanç varsa serisinde (401–483) işlenmeyen veya kısmi kalan kalemler."
)

grp_b_table = doc.add_table(rows=4, cols=5)
grp_b_table.style = "Table Grid"
grp_b_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate([Cm(1.5), Cm(5), Cm(3), Cm(2), Cm(4)]):
    for row in grp_b_table.rows:
        row.cells[i].width = w

header_row(grp_b_table, ["Kod", "ÇK Adı", "Mevzuat", "Beyanname Kodu", "Notlar"])

grp_b_data = [
    ("ÇK-B1", "Finans/Bankacılık Sektörü Temettü İndirimi", "KVK Md. 10/1-j", "484", "XML'de kayıt mevcut; YAML henüz oluşturulmadı."),
    ("ÇK-B2", "Türkiye Varlık Fonu Bağışı", "TVF Kanunu", "—", "Yeni mevzuatla eklendi; beyanname kodu teyit edilecek."),
    ("ÇK-B3", "Kooperatif Risturn Tamamlama", "KVK Md. 5/1-i", "401 mevcut", "Tüketim / üretim / kredi kooperatifi ayrımına göre alt ÇK'lar gerekebilir."),
]

for i, row_data in enumerate(grp_b_data, start=1):
    data_row(grp_b_table, i, list(row_data))

doc.add_paragraph()

# ─── Bölüm 4: GRUP XML (389–400 + 484) ───────────────────────────────────────
add_heading(doc, "4. Grup XML — GIB Beyanname XML'inde Olan, YAML Kataloğunda Olmayan Kalemler", 1)
doc.add_paragraph(
    "Aşağıdaki kalemler GIB e-beyanname kod listesinde (KURUMLAR_29_Kodlar.xml) yer almakta; "
    "ancak Declaro kalem kataloğunda (kalemler/*.yaml) karşılıkları bulunmamaktadır. "
    "Tümü KVK Madde 5/1-d kapsamında portföy işletmeciliği istisna kalemleridir; "
    "taşınmaz dahil/hariç ayrımına göre çiftler halinde tanımlanmıştır."
)

grp_xml_table = doc.add_table(rows=14, cols=5)
grp_xml_table.style = "Table Grid"
grp_xml_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate([Cm(1.5), Cm(1.8), Cm(5), Cm(3), Cm(4.2)]):
    for row in grp_xml_table.rows:
        row.cells[i].width = w

header_row(grp_xml_table, ["Kod", "Beyanname No", "Kalem Adı", "Mevzuat", "Notlar"])

grp_xml_data = [
    # ── Zarar Olsa Dahi grubu (389–400)
    ("ÇK-X01", "389", "MK Yatırım Fonu/Ortaklığı — Portföy İşletmeciliği (Taşınmaz Hariç)",
     "KVK Md. 5/1-d-1", "298 GSYF/302 Portföy ile aynı ana gruba giriyor; ancak MKYF'ye özgü taşınmaz hariç ayrımı."),
    ("ÇK-X02", "390", "MK Yatırım Fonu/Ortaklığı — Portföy İşletmeciliği (Taşınmaz Dahil)",
     "KVK Md. 5/1-d-1", "389 ile çift; taşınmazdan elde edilen kısım için ayrı satır."),
    ("ÇK-X03", "391", "Araştırma Altyapılarının Ar-Ge ve Yenilik Faaliyeti Kazancı İstisnası",
     "6550 s. Kanun", "Üniversite bünyesindeki araştırma altyapıları için özel istisna; zarar_olsa_dahi."),
    ("ÇK-X04", "392", "Altın/Kıymetli Maden Fonları — Portföy İşletmeciliği (Taşınmaz Hariç)",
     "KVK Md. 5/1-d-2", "Borsa'da işlem gören altın/kıymetli maden fonları; taşınmaz hariç kısım."),
    ("ÇK-X05", "393", "Altın/Kıymetli Maden Fonları — Portföy İşletmeciliği (Taşınmaz Dahil)",
     "KVK Md. 5/1-d-2", "392 ile çift; taşınmazdan elde edilen kısım."),
    ("ÇK-X06", "394", "Girişim Sermayesi Yatırım Fonu/Ortaklığı (Taşınmaz Hariç)",
     "KVK Md. 5/1-d-3", "413 VUK325A GİSEF ile ilişkili; ancak bu satır doğrudan kazanç istisnası."),
    ("ÇK-X07", "395", "Girişim Sermayesi Yatırım Fonu/Ortaklığı (Taşınmaz Dahil)",
     "KVK Md. 5/1-d-3", "394 ile çift."),
    ("ÇK-X08", "396", "Gayrimenkul Yatırım Fonu/Ortaklığı Kazancı (Taşınmaz Hariç)",
     "KVK Md. 5/1-d-4", "GYF/GYO portföy kazancı; taşınmaz dışı kısım."),
    ("ÇK-X09", "397", "Gayrimenkul Yatırım Fonu/Ortaklığı Kazancı (Taşınmaz Dahil)",
     "KVK Md. 5/1-d-4", "396 ile çift; taşınmazdan elde edilen kısım."),
    ("ÇK-X10", "398", "Emeklilik Yatırım Fonu Kazancı (Taşınmaz Hariç)",
     "KVK Md. 5/1-d-5", "EYF'lerin portföy işletmeciliği; taşınmaz dışı."),
    ("ÇK-X11", "399", "Emeklilik Yatırım Fonu Kazancı (Taşınmaz Dahil)",
     "KVK Md. 5/1-d-5", "398 ile çift; taşınmazdan elde edilen kısım."),
    ("ÇK-X12", "400", "Konut/Varlık Finansman Fonu Kazancı",
     "KVK Md. 5/1-d", "KFF ve VFF için özel portföy kazancı istisnası."),
    # ── Kazanç Varsa grubu (484)
    ("ÇK-X13", "484", "TENMAK Bağışı / İndirimi",
     "2690 s. Kanun", "Türkiye Enerji, Nükleer ve Maden Araştırma Kurumu'na yapılan bağış; kazanç_varsa."),
]

for i, row_data in enumerate(grp_xml_data, start=1):
    # Kazanç varsa kalemine (484) farklı renk
    bg = "FFF0E6" if row_data[1] == "484" else None
    data_row(grp_xml_table, i, list(row_data), bg=bg)

add_info_box(
    doc,
    "⚑  389–400 arası kalemler 'zarar_olsa_dahi' bölümüne aittir. "
    "484 numaralı TENMAK kalemi 'kazanc_varsa' grubundadır (turuncu satır). "
    "Her çift (taşınmaz dahil/hariç), tek YAML dosyasında seçenek alanı ile çözülebilir.",
)

doc.add_paragraph()

# ─── Bölüm 5: GRUP C ──────────────────────────────────────────────────────────
add_heading(doc, "5. Grup C — Zarar Olsa Dahi Serisi Diğer Eksikler", 1)
doc.add_paragraph(
    "Zarar olsa dahi serisinde (297–387) tespit edilen diğer boşluklar."
)

grp_c_table = doc.add_table(rows=3, cols=5)
grp_c_table.style = "Table Grid"
grp_c_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate([Cm(1.5), Cm(5), Cm(3), Cm(2), Cm(4)]):
    for row in grp_c_table.rows:
        row.cells[i].width = w

header_row(grp_c_table, ["Kod", "ÇK Adı", "Mevzuat", "Beyanname Kodu", "Notlar"])

grp_c_data = [
    ("ÇK-C1", "Enflasyon Düzeltmesi İstisnası", "VUK 298/Ç", "388", "2023–2024 enflasyon düzeltmesinden kaynaklanan kazanç istisnası; özel hesaplama gerektirir."),
    ("ÇK-C2", "Finansman Fonu Girdi Sayfası", "VUK Md. 390", "—", "Pipeline adım 2'de parametre mevcut; ancak kullanıcı arayüzü/ÇK tanımlanmamış."),
]

for i, row_data in enumerate(grp_c_data, start=1):
    data_row(grp_c_table, i, list(row_data))

doc.add_paragraph()

# ─── Bölüm 6: GRUP D ──────────────────────────────────────────────────────────
add_heading(doc, "6. Grup D — Yapısal İyileştirme / Görünürlük", 1)
doc.add_paragraph(
    "Hesaplama altyapısı kısmen mevcut, ancak kullanıcıya görünür ÇK / özet sayfası eksik."
)

grp_d_table = doc.add_table(rows=3, cols=4)
grp_d_table.style = "Table Grid"
grp_d_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate([Cm(1.5), Cm(5.5), Cm(3), Cm(5)]):
    for row in grp_d_table.rows:
        row.cells[i].width = w

header_row(grp_d_table, ["Kod", "ÇK / Modül Adı", "Mevzuat", "Notlar"])

grp_d_data = [
    ("ÇK-D1", "YİAKV Hesap Özeti Sayfası", "KVK Md. 32/C",
     "docs/YİAKV_Calisma_Kagidi_KVK_32C_GIB2026_Rev4.xlsx VAR. Pipeline hesaplıyor; "
     "ancak kullanıcı hangi kalemlerin YİAKV matrahından düşüldüğünü görmüyor."),
    ("ÇK-D2", "Devreden Ar-Ge / Tasarım İndirimi Takibi", "KVK 10/1-a + 5746",
     "409 ve 410 YAML'larında 'devreden_sonraki_doneme' hesaplanıyor; "
     "dönemler arası devir takip modülü tanımlanmamış."),
]

for i, row_data in enumerate(grp_d_data, start=1):
    data_row(grp_d_table, i, list(row_data))

doc.add_paragraph()

# ─── Bölüm 7: Öncelik Sırası ─────────────────────────────────────────────────
add_heading(doc, "7. Önerilen Kodlama / ÇK Hazırlık Sırası", 1)
doc.add_paragraph(
    "Aşağıdaki sıra; etkilenen mükellef sayısı, pipeline bütünlüğü ve mevcut Excel ÇK varlığına göre belirlenmiştir."
)

prio_table = doc.add_table(rows=16, cols=4)
prio_table.style = "Table Grid"
prio_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, w in enumerate([Cm(1.2), Cm(2), Cm(5.5), Cm(6.3)]):
    for row in prio_table.rows:
        row.cells[i].width = w

header_row(prio_table, ["Sıra", "ÇK Kodu", "ÇK Adı", "Gerekçe"])

prio_data = [
    ("1", "ÇK-A1", "Geçmiş Yıl Zararları Mahsubu", "Neredeyse her mükellef kullanır; pipeline adım 5 boş."),
    ("2", "ÇK-A2", "KVK 32/A Tek YTB", "YTB sahibi kurumlar için zorunlu; Excel ÇK mevcut."),
    ("3", "ÇK-A3", "KVK 32/A Çok YTB", "Excel ÇK docs/ altında hazır; sadece YAML dönüşümü gerekiyor."),
    ("4", "ÇK-A4", "KVK 32/7 KOBİ İndirimi", "2022+ tüm sanayi siciline kayıtlı üretim işletmeleri."),
    ("5", "ÇK-A5", "KVK 32/8 İhracat İndirimi", "İhracatçı kurumlar için kritik."),
    ("6", "ÇK-X01–12", "KVK 5/1-d Fon/Ortaklık İstisnaları (389–400)", "GIB XML'de mevcut; YAML yok. 12 kalem — taşınmaz dahil/hariç ayrımı. Fonlarla çalışan mükellefler için zorunlu."),
    ("7", "ÇK-X13", "TENMAK Bağışı (484)", "GIB XML'de mevcut; YAML yok. Kazanç varsa grubuna eklenecek."),
    ("8", "ÇK-A6", "Geçici Vergi Mahsubu", "Pipeline adım 12; ödenecek vergiyi doğrudan etkiler."),
    ("9", "ÇK-A7", "Yurt Dışı Vergi Mahsubu", "Pipeline adım 12; uluslararası işlemler."),
    ("10", "ÇK-A8", "Tevkifat Mahsubu", "Pipeline adım 12; stopaj takibi."),
    ("11", "ÇK-X03", "6550 s. K. Araştırma Altyapısı İstisnası (391)", "GIB XML'de mevcut; üniversite bünyesindeki araştırma altyapıları."),
    ("12", "ÇK-C1", "Enflasyon Düzeltmesi İstisnası", "2023–2024 özel; zarar grubuna ek kalem."),
    ("13", "ÇK-C2", "Finansman Fonu Girdi Sayfası", "UI/UX düzeltme; kullanıcı girişi tanımlanacak."),
    ("14", "ÇK-D1", "YİAKV Özet Sayfası", "Görünürlük iyileştirmesi; hesaplama altyapısı mevcut."),
    ("15", "ÇK-D2", "Devreden Ar-Ge Takibi", "Dönemler arası devir modülü; Ar-Ge yoğun mükellefleri etkiler."),
]

for i, row_data in enumerate(prio_data, start=1):
    # İlk 7'yi vurgula
    bg = "FFF3CD" if i <= 7 else None
    data_row(prio_table, i, list(row_data), bg=bg)

doc.add_paragraph()

# ─── Dipnot ───────────────────────────────────────────────────────────────────
p_note = doc.add_paragraph()
run_note = p_note.add_run(
    "Not: Bu belge Declaro teknik altyapısının (YAML kalem kataloğu + Python pipeline) "
    "GIB e-beyanname kodu listesi ve mevzuat kapsamı ile karşılaştırılması sonucunda "
    "oluşturulmuştur. Sarı satırlar en yüksek öncelikli kalemleri göstermektedir."
)
run_note.italic = True
run_note.font.size = Pt(8)
run_note.font.color.rgb = RGBColor.from_string("666666")

# ─── Kaydet ───────────────────────────────────────────────────────────────────
out_path = "/home/ziyahan/declaro/docs/Declaro_Gap_Analizi_CK_Listesi.docx"
doc.save(out_path)
print(f"✓ Belge oluşturuldu: {out_path}")
